from __future__ import annotations

import argparse
import json
import os
import random
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any

import pyarrow as pa
import pyarrow.parquet as pq


DEFAULT_REPROMPT_TEMPLATE = """{prompt}

Previous attempt:
{solution}

Harvey LAB rubric feedback:
{feedback}

Use the feedback to produce a corrected legal work product. Preserve the requested deliverable structure."""


@dataclass(frozen=True)
class TaskSpec:
    task_id: str
    task_path: Path
    practice_area: str
    title: str
    work_type: str
    instructions: str
    deliverables: list[str]
    criteria: list[dict[str, Any]]


def discover_tasks(bench_root: Path) -> list[TaskSpec]:
    tasks: list[TaskSpec] = []
    for task_json in sorted((bench_root / "tasks").rglob("task.json")):
        raw = json.loads(task_json.read_text(encoding="utf-8"))
        task_id = task_json.parent.relative_to(bench_root / "tasks").as_posix()
        parts = task_id.split("/")
        deliverables = sorted(
            {
                str(item)
                for criterion in raw.get("criteria", [])
                for item in criterion.get("deliverables", [])
            }
        )
        if not deliverables:
            deliverables = sorted(str(v) for v in raw.get("deliverables", {}).values())
        tasks.append(
            TaskSpec(
                task_id=task_id,
                task_path=task_json.parent,
                practice_area=parts[0],
                title=str(raw.get("title", "")),
                work_type=str(raw.get("work_type", "")),
                instructions=str(raw.get("instructions", "")),
                deliverables=deliverables or ["output.md"],
                criteria=list(raw.get("criteria", [])),
            )
        )
    return tasks


def stratified_split(
    tasks: list[TaskSpec],
    *,
    train_tasks: int,
    eval_preview_tasks: int,
    seed: int,
) -> tuple[list[TaskSpec], list[TaskSpec], list[TaskSpec]]:
    if train_tasks <= 0:
        raise ValueError("train_tasks must be positive")
    rng = random.Random(seed)
    buckets: dict[str, list[TaskSpec]] = {}
    for task in tasks:
        buckets.setdefault(task.practice_area, []).append(task)
    for bucket in buckets.values():
        rng.shuffle(bucket)

    ordered: list[TaskSpec] = []
    while len(ordered) < len(tasks):
        progressed = False
        for area in sorted(buckets):
            if buckets[area]:
                ordered.append(buckets[area].pop())
                progressed = True
        if not progressed:
            break

    train = ordered[: min(train_tasks, len(ordered))]
    train_ids = {task.task_id for task in train}
    heldout = [task for task in ordered if task.task_id not in train_ids]
    eval_preview = heldout[: max(eval_preview_tasks, 0)]
    return train, heldout, eval_preview


def _trim(text: str, limit: int) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    if len(text) <= limit:
        return text
    return text[:limit].rsplit("\n", 1)[0] + "\n[truncated]"


def read_document(path: Path, limit: int) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            from docx import Document

            doc = Document(str(path))
            return _trim("\n".join(p.text for p in doc.paragraphs if p.text.strip()), limit)
        if suffix == ".xlsx":
            from openpyxl import load_workbook

            wb = load_workbook(path, read_only=True, data_only=True)
            parts: list[str] = []
            for sheet in wb.worksheets[:4]:
                parts.append(f"=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(max_row=40, max_col=12, values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(values):
                        parts.append("\t".join(values))
            return _trim("\n".join(parts), limit)
        if suffix == ".eml":
            msg = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
            body = msg.get_body(preferencelist=("plain", "html"))
            content = body.get_content() if body else path.read_text(encoding="utf-8", errors="replace")
            return _trim(content, limit)
        if suffix in {".txt", ".md", ".json", ".csv"}:
            return _trim(path.read_text(encoding="utf-8", errors="replace"), limit)
    except Exception as exc:
        return f"(error reading {path.name}: {exc})"
    return f"(document available but not extracted: {path.name})"


def build_prompt(task: TaskSpec, *, doc_chars: int, max_prompt_chars: int, include_rubric: bool) -> str:
    sections = [
        "You are completing a Harvey LAB legal-agent benchmark task.",
        f"Task ID: {task.task_id}",
        f"Title: {task.title}",
        f"Work type: {task.work_type}",
        f"Instructions: {task.instructions}",
        "Expected deliverables:",
        "\n".join(f"- {name}" for name in task.deliverables),
        "Source document excerpts:",
    ]
    docs_dir = task.task_path / "documents"
    if docs_dir.exists():
        for doc_path in sorted(path for path in docs_dir.rglob("*") if path.is_file()):
            sections.append(f"\n### {doc_path.relative_to(docs_dir)}\n{read_document(doc_path, doc_chars)}")
    if include_rubric:
        rubric = []
        for criterion in task.criteria:
            rubric.append(
                f"- {criterion.get('id', '')}: {criterion.get('title', '')}\n"
                f"  {criterion.get('match_criteria', '')}"
            )
        sections.extend(["Rubric criteria for training feedback construction:", "\n".join(rubric)])
    sections.append("Draft the requested legal work product. Use clear headings and cite source documents when possible.")
    return _trim("\n\n".join(sections), max_prompt_chars)


def fallback_attempt(task: TaskSpec) -> str:
    missed = "\n".join(
        f"- TODO {criterion.get('id', '')}: {criterion.get('title', '')}"
        for criterion in task.criteria[:12]
    )
    return (
        f"# Draft work product for {task.title}\n\n"
        "This is an intentionally sparse baseline attempt generated for the SDPO smoke path.\n\n"
        "Open rubric-sensitive issues to address:\n"
        f"{missed}\n"
    )


def generate_attempts(
    prompts: list[str],
    *,
    model: str,
    tensor_parallel_size: int,
    max_model_len: int,
    max_new_tokens: int,
    skip_generation: bool,
) -> list[str]:
    if skip_generation:
        return []
    from vllm import LLM, SamplingParams

    llm = LLM(
        model=model,
        trust_remote_code=True,
        tensor_parallel_size=tensor_parallel_size,
        max_model_len=max_model_len,
        gpu_memory_utilization=0.9,
    )
    outputs = llm.generate(
        prompts,
        SamplingParams(temperature=0.6, top_p=0.95, max_tokens=max_new_tokens),
    )
    return [item.outputs[0].text.strip() for item in outputs]


def write_deliverable(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        from docx import Document

        doc = Document()
        for paragraph in content.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(str(path))
        return
    if suffix == ".xlsx":
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Draft"
        for row_idx, line in enumerate(content.split("\n")[:200], start=1):
            ws.cell(row=row_idx, column=1, value=line)
        wb.save(str(path))
        return
    path.write_text(content, encoding="utf-8")


def write_run_outputs(bench_root: Path, task: TaskSpec, attempt: str, run_id: str) -> Path:
    output_dir = bench_root / "results" / run_id / "output"
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    for deliverable in task.deliverables:
        write_deliverable(output_dir / deliverable, attempt)
    metrics = {
        "input_tokens": 0,
        "output_tokens": 0,
        "wall_clock_seconds": 0,
        "documents_read": len(list((task.task_path / "documents").rglob("*"))) if (task.task_path / "documents").exists() else 0,
        "total_vdr_files": len(list((task.task_path / "documents").rglob("*"))) if (task.task_path / "documents").exists() else 0,
    }
    (bench_root / "results" / run_id / "metrics.json").write_text(json.dumps(metrics, indent=2))
    return output_dir


def run_harvey_eval(
    bench_root: Path,
    *,
    run_id: str,
    task_id: str,
    judge_model: str,
    judge_parallel: int,
) -> dict[str, Any]:
    cmd = [
        sys.executable,
        "-m",
        "evaluation.run_eval",
        "--run-id",
        run_id,
        "--task",
        task_id,
        "--judge-model",
        judge_model,
        "--parallel",
        str(judge_parallel),
    ]
    subprocess.run(cmd, cwd=bench_root, check=True)
    return json.loads((bench_root / "results" / run_id / "scores.json").read_text(encoding="utf-8"))


def reward_terms(task: TaskSpec, limit: int = 48) -> list[str]:
    raw = " ".join(
        f"{criterion.get('title', '')} {criterion.get('match_criteria', '')}"
        for criterion in task.criteria
    )
    quoted = re.findall(r"[\"'`]([^\"'`]{3,80})[\"'`]", raw)
    numbers = re.findall(r"\b\d+(?:\.\d+)?%?x?\b", raw)
    caps = re.findall(r"\b[A-Z][A-Za-z0-9&/-]*(?:\s+[A-Z][A-Za-z0-9&/-]*){0,5}\b", raw)
    terms: list[str] = []
    for value in quoted + numbers + caps:
        value = re.sub(r"\s+", " ", value).strip()
        if len(value) >= 3 and value.lower() not in {"pass", "fail", "agent", "issue"} and value not in terms:
            terms.append(value)
        if len(terms) >= limit:
            break
    return terms


def feedback_from_scores(task: TaskSpec, scores: dict[str, Any] | None, limit: int = 3200) -> str:
    if not scores:
        return _trim("\n".join(
            f"- {criterion.get('id', '')} {criterion.get('title', '')}: {criterion.get('match_criteria', '')}"
            for criterion in task.criteria[:12]
        ), limit)
    lines = [scores.get("summary", "Harvey LAB rubric feedback")]
    for result in scores.get("criteria_results", []):
        verdict = str(result.get("verdict", "fail")).upper()
        lines.append(
            f"- {verdict} {result.get('id', '')}: {result.get('title', '')}\n"
            f"  Judge reasoning: {result.get('reasoning', '')}"
        )
        if verdict == "FAIL":
            match = next((c.get("match_criteria", "") for c in task.criteria if c.get("id") == result.get("id")), "")
            if match:
                lines.append(f"  Required standard: {match}")
    return _trim("\n".join(lines), limit)


def make_sdpo_record(
    task: TaskSpec,
    *,
    split: str,
    prompt: str,
    previous_attempt: str,
    feedback: str,
    scores: dict[str, Any] | None,
    reprompt_template: str,
    include_feedback: bool,
) -> dict[str, Any]:
    final_prompt = (
        reprompt_template.format(prompt=prompt, solution=previous_attempt, feedback=feedback)
        if include_feedback
        else prompt
    )
    return {
        "data_source": "harvey/lab",
        "prompt": [{"role": "user", "content": final_prompt}],
        "ability": "legal_agent",
        "reward_model": {"style": "rule", "ground_truth": "harvey_lab_rubric"},
        "extra_info": {
            "split": split,
            "task_id": task.task_id,
            "practice_area": task.practice_area,
            "work_type": task.work_type,
            "original_prompt": prompt,
            "answer": "harvey_lab_rubric",
            "feedback_raw": feedback if include_feedback else "",
            "previous_attempt": previous_attempt if include_feedback else "",
            "harvey_scores_json": json.dumps(scores or {}, sort_keys=True),
            "reward_terms": reward_terms(task),
            "sdpo_bridge": True,
        },
    }


def write_parquet(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pq.write_table(pa.Table.from_pylist(rows), path)


def prepare(args: argparse.Namespace) -> None:
    bench_root = Path(args.bench_root)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    tasks = discover_tasks(bench_root)
    train, heldout, eval_preview = stratified_split(
        tasks,
        train_tasks=args.train_tasks,
        eval_preview_tasks=args.eval_preview_tasks,
        seed=args.seed,
    )
    train_prompt_tasks = train
    eval_prompt_tasks = eval_preview or heldout[:1]
    all_generated_tasks = train_prompt_tasks + eval_prompt_tasks
    prompts = [
        build_prompt(
            task,
            doc_chars=args.doc_chars,
            max_prompt_chars=args.max_prompt_chars,
            include_rubric=False,
        )
        for task in all_generated_tasks
    ]
    generated = generate_attempts(
        prompts,
        model=args.model,
        tensor_parallel_size=args.tensor_parallel_size,
        max_model_len=args.max_model_len,
        max_new_tokens=args.max_new_tokens,
        skip_generation=args.skip_generation,
    )
    attempts: dict[str, str] = {}
    for idx, task in enumerate(all_generated_tasks):
        attempts[task.task_id] = generated[idx] if generated else fallback_attempt(task)

    scores_by_task: dict[str, dict[str, Any]] = {}
    run_manifest: list[dict[str, Any]] = []
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    for task in all_generated_tasks:
        run_id = f"harvey_sdpo/{task.task_id}/qwen36_35b_a3b/{stamp}"
        write_run_outputs(bench_root, task, attempts[task.task_id], run_id)
        scores = None
        if not args.skip_judge_eval:
            scores = run_harvey_eval(
                bench_root,
                run_id=run_id,
                task_id=task.task_id,
                judge_model=args.judge_model,
                judge_parallel=args.judge_parallel,
            )
            scores_by_task[task.task_id] = scores
            scores_out = out_dir / "scores" / task.task_id / "scores.json"
            scores_out.parent.mkdir(parents=True, exist_ok=True)
            scores_out.write_text(json.dumps(scores, indent=2), encoding="utf-8")
        run_manifest.append(
            {
                "task_id": task.task_id,
                "split": "train" if task in train_prompt_tasks else "eval_preview",
                "run_id": run_id,
                "score": None if not scores else scores.get("score"),
                "all_pass": None if not scores else scores.get("all_pass"),
            }
        )

    train_rows = []
    for task in train_prompt_tasks:
        prompt = build_prompt(task, doc_chars=args.doc_chars, max_prompt_chars=args.max_prompt_chars, include_rubric=False)
        scores = scores_by_task.get(task.task_id)
        train_rows.append(
            make_sdpo_record(
                task,
                split="train",
                prompt=prompt,
                previous_attempt=attempts[task.task_id],
                feedback=feedback_from_scores(task, scores),
                scores=scores,
                reprompt_template=DEFAULT_REPROMPT_TEMPLATE,
                include_feedback=True,
            )
        )

    val_rows = []
    for task in eval_prompt_tasks:
        prompt = build_prompt(task, doc_chars=args.doc_chars, max_prompt_chars=args.max_prompt_chars, include_rubric=False)
        val_rows.append(
            make_sdpo_record(
                task,
                split="test",
                prompt=prompt,
                previous_attempt=attempts[task.task_id],
                feedback="Held-out Harvey LAB task. Feedback is intentionally withheld from the validation prompt.",
                scores=scores_by_task.get(task.task_id),
                reprompt_template=DEFAULT_REPROMPT_TEMPLATE,
                include_feedback=False,
            )
        )

    write_parquet(out_dir / "train.parquet", train_rows)
    write_parquet(out_dir / "test.parquet", val_rows)
    manifest = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "model": args.model,
        "train_tasks": [task.task_id for task in train],
        "heldout_eval_tasks": [task.task_id for task in heldout],
        "eval_preview_tasks": [task.task_id for task in eval_prompt_tasks],
        "runs": run_manifest,
        "paths": {
            "train_parquet": str(out_dir / "train.parquet"),
            "test_parquet": str(out_dir / "test.parquet"),
        },
        "notes": [
            "Training rows include rubric-derived feedback from generated attempts.",
            "Held-out eval tasks are recorded in the manifest and are not used as feedback-conditioned train rows.",
        ],
    }
    (out_dir / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Prepare Harvey LAB SDPO data on Modal.")
    parser.add_argument("--bench-root", required=True)
    parser.add_argument("--out-dir", required=True)
    parser.add_argument("--model", required=True)
    parser.add_argument("--train-tasks", type=int, default=8)
    parser.add_argument("--eval-preview-tasks", type=int, default=2)
    parser.add_argument("--seed", type=int, default=17)
    parser.add_argument("--doc-chars", type=int, default=1200)
    parser.add_argument("--max-prompt-chars", type=int, default=9000)
    parser.add_argument("--tensor-parallel-size", type=int, default=4)
    parser.add_argument("--max-model-len", type=int, default=12288)
    parser.add_argument("--max-new-tokens", type=int, default=1024)
    parser.add_argument("--judge-model", default="claude-sonnet-4-6")
    parser.add_argument("--judge-parallel", type=int, default=2)
    parser.add_argument("--skip-generation", action="store_true")
    parser.add_argument("--skip-judge-eval", action="store_true")
    return parser


def main() -> None:
    prepare(build_parser().parse_args())


if __name__ == "__main__":
    main()
